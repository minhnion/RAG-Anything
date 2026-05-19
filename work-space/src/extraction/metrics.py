import hashlib
import math
import re
from pathlib import Path
from statistics import median
from typing import Any, Dict, Iterable, List, Optional

DOCX_CHARS_PER_PAGE = 1800


def _count_tokens(text: str) -> int:
    if not text:
        return 0
    try:
        import tiktoken

        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(str(text)))
    except Exception:
        return 0


_PAGE_NUMBER_RE = re.compile(r"^\s*(?:page\s*)?\d+(?:\s*/\s*\d+)?\s*$", re.IGNORECASE)
_COORDINATE_RE = re.compile(r"\b(?:bbox|x0|y0|x1|y1|left|top|right|bottom)\b", re.IGNORECASE)
_ALNUM_RE = re.compile(r"[A-Za-z0-9]")


def _parse_markdown_table(table_str: str) -> Dict[str, int]:
    rows = 0
    cols = 0
    cells = 0
    if not table_str:
        return {"rows": 0, "cols": 0, "cells": 0}

    lines = [l.strip() for l in str(table_str).splitlines() if l.strip()]
    for line in lines:
        if "|" not in line:
            continue
        if set(line.replace("|", "").replace(":", "").replace("-", "").replace(" ", "")) == set():
            continue
        parts = [p.strip() for p in line.split("|") if p.strip()]
        if not parts:
            continue
        rows += 1
        cols = max(cols, len(parts))
        cells += len(parts)

    return {"rows": rows, "cols": cols, "cells": cells}


def _parse_html_table(table_str: str) -> Dict[str, int]:
    rows = 0
    cols = 0
    cells = 0
    if not table_str:
        return {"rows": 0, "cols": 0, "cells": 0}

    tr_blocks = re.findall(r"<tr[^>]*>(.*?)</tr>", str(table_str), flags=re.IGNORECASE | re.DOTALL)
    for tr in tr_blocks:
        cell_blocks = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", tr, flags=re.IGNORECASE | re.DOTALL)
        cleaned_cells = [
            re.sub(r"<[^>]+>", "", c).strip() for c in cell_blocks if re.sub(r"<[^>]+>", "", c).strip()
        ]
        if not cleaned_cells:
            continue
        rows += 1
        cols = max(cols, len(cleaned_cells))
        cells += len(cleaned_cells)

    return {"rows": rows, "cols": cols, "cells": cells}


def _parse_table_body(table_body: Any) -> Dict[str, int]:
    if table_body is None:
        return {"rows": 0, "cols": 0, "cells": 0}

    if isinstance(table_body, list):
        rows = len(table_body)
        cols = 0
        cells = 0
        for row in table_body:
            if isinstance(row, list):
                cols = max(cols, len(row))
                cells += len(row)
            elif isinstance(row, dict):
                cols = max(cols, len(row))
                cells += len(row)
            else:
                cells += 1
        return {"rows": rows, "cols": cols, "cells": cells}

    if isinstance(table_body, dict):
        num_rows = int(table_body.get("num_rows", 0) or 0)
        num_cols = int(table_body.get("num_cols", 0) or 0)
        grid = table_body.get("grid")

        if isinstance(grid, list) and grid:
            rows = 0
            cols = 0
            cells = 0
            for row in grid:
                if not isinstance(row, list):
                    continue
                non_empty = 0
                for cell in row:
                    if isinstance(cell, dict):
                        text = str(cell.get("text", "")).strip()
                    else:
                        text = str(cell).strip()
                    if text:
                        non_empty += 1
                if non_empty > 0:
                    rows += 1
                    cols = max(cols, len(row))
                    cells += non_empty
            if rows > 0:
                return {"rows": rows, "cols": cols, "cells": cells}

        table_cells = table_body.get("table_cells")
        if isinstance(table_cells, list) and table_cells:
            if num_rows > 0 and num_cols > 0:
                return {"rows": num_rows, "cols": num_cols, "cells": len(table_cells)}
            return {"rows": 0, "cols": 0, "cells": len(table_cells)}

        if num_rows > 0 or num_cols > 0:
            return {"rows": num_rows, "cols": num_cols, "cells": max(num_rows * num_cols, 0)}

        return {"rows": 0, "cols": 0, "cells": 0}

    if isinstance(table_body, str):
        if "<table" in table_body.lower():
            return _parse_html_table(table_body)
        return _parse_markdown_table(table_body)

    return {"rows": 0, "cols": 0, "cells": 0}


def _is_valid_image_file(path: str) -> bool:
    try:
        with open(path, "rb") as f:
            head = f.read(16)
    except Exception:
        return False

    known = [
        b"\x89PNG\r\n\x1a\n",
        b"\xff\xd8\xff",
        b"GIF87a",
        b"GIF89a",
        b"BM",
        b"II*\x00",
        b"MM\x00*",
    ]
    if any(head.startswith(sig) for sig in known):
        return True
    if head.startswith(b"RIFF") and head[8:12] == b"WEBP":
        return True

    try:
        from PIL import Image

        with Image.open(path) as img:
            img.verify()
        return True
    except Exception:
        return False


def _estimate_source_pages(content_list: List[Dict[str, Any]]) -> int:
    page_ids = set()
    for item in content_list:
        if not isinstance(item, dict):
            continue
        page_idx = item.get("page_idx")
        if isinstance(page_idx, int) and page_idx >= 0:
            page_ids.add(page_idx)
    if page_ids:
        return max(len(page_ids), max(page_ids) + 1)
    return 1 if content_list else 0


def get_source_page_count(file_path: str | Path) -> int:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp", ".gif"}:
        return 1

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            with open(path, "rb") as f:
                return len(PdfReader(f).pages)
        except Exception:
            try:
                import fitz

                doc = fitz.open(path)
                try:
                    return int(doc.page_count)
                finally:
                    doc.close()
            except Exception:
                return 0

    if suffix == ".pptx":
        try:
            from pptx import Presentation

            return len(Presentation(str(path)).slides)
        except Exception:
            return 0

    if suffix == ".docx":
        # docx has no fixed page concept; Word computes pages only at render time.
        # Heuristic: total visible char count / DOCX_CHARS_PER_PAGE (~500 words/page).
        try:
            from docx import Document

            doc = Document(str(path))
            total_chars = sum(len(p.text or "") for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        total_chars += len(cell.text or "")
            if total_chars <= 0:
                return 0
            return max(1, math.ceil(total_chars / DOCX_CHARS_PER_PAGE))
        except Exception:
            return 0

    return 0


def _is_symbol_heavy(line: str) -> bool:
    stripped = line.strip()
    if len(stripped) < 6:
        return False
    non_space = [ch for ch in stripped if not ch.isspace()]
    if not non_space:
        return False
    symbol_count = sum(1 for ch in non_space if not ch.isalnum())
    return (symbol_count / len(non_space)) >= 0.65 and not _ALNUM_RE.search(stripped)


def _collect_repeated_short_lines(items: List[Dict[str, Any]]) -> set[str]:
    line_pages: Dict[str, set[int]] = {}
    for item in items:
        if not isinstance(item, dict) or item.get("type") != "text":
            continue
        page_idx = item.get("page_idx", 0)
        if not isinstance(page_idx, int):
            page_idx = 0
        text = str(item.get("text", "") or "")
        for line in text.splitlines():
            normalized = re.sub(r"\s+", " ", line).strip()
            if len(normalized) < 3 or len(normalized) > 80:
                continue
            if len(normalized.split()) > 12:
                continue
            line_pages.setdefault(normalized, set()).add(page_idx)
    return {line for line, pages in line_pages.items() if len(pages) >= 3}


def _estimate_noise_ratio(items: List[Dict[str, Any]], text_chars: int) -> float:
    if text_chars <= 0:
        return 0.0

    repeated_short_lines = _collect_repeated_short_lines(items)
    junk_chars = 0
    for item in items:
        if not isinstance(item, dict) or item.get("type") != "text":
            continue
        text = str(item.get("text", "") or "")
        for line in text.splitlines() or [text]:
            normalized = re.sub(r"\s+", " ", line).strip()
            if not normalized:
                continue
            is_junk = False
            if _PAGE_NUMBER_RE.fullmatch(normalized):
                is_junk = True
            elif _COORDINATE_RE.search(normalized):
                is_junk = True
            elif normalized in repeated_short_lines:
                is_junk = True
            elif _is_symbol_heavy(normalized):
                is_junk = True
            if is_junk:
                junk_chars += len(normalized)

    return min(junk_chars / max(text_chars, 1), 1.0)


def _safe_rate(numerator: float, denominator: float, scale: float = 1.0) -> float:
    if denominator <= 0:
        return 0.0
    return (numerator / denominator) * scale


def _format_modality_profile(metrics: Dict[str, Any]) -> str:
    return (
        f"text={metrics['text_blocks_per_page']:.1f}/page | "
        f"table={metrics['tables_per_100_pages']:.1f}/100p | "
        f"figure={metrics['figures_per_100_pages']:.1f}/100p | "
        f"equation={metrics['equations_per_100_pages']:.1f}/100p"
    )


def compute_extract_metrics(
    content_list: List[Dict[str, Any]],
    source_pages_override: Optional[int] = None,
) -> Dict[str, Any]:
    metrics = {
        "total_blocks": 0,
        "source_pages": 0,
        "text_blocks": 0,
        "empty_text_blocks": 0,
        "text_chars": 0,
        "text_tokens": 0,
        "image_blocks": 0,
        "image_files_exist": 0,
        "image_files_missing": 0,
        "table_blocks": 0,
        "table_rows": 0,
        "table_cells": 0,
        "equation_blocks": 0,
        "text_md5": "",
        "noise_ratio": 0.0,
        "tokens_per_page": 0.0,
        "text_blocks_per_page": 0.0,
        "tables_per_100_pages": 0.0,
        "figures_per_100_pages": 0.0,
        "equations_per_100_pages": 0.0,
        "modality_coverage_profile": "",
    }

    text_parts: List[str] = []
    metrics["total_blocks"] = len(content_list)
    estimated_pages = _estimate_source_pages(content_list)
    metrics["source_pages"] = int(source_pages_override or 0) if int(source_pages_override or 0) > 0 else estimated_pages

    for item in content_list:
        if not isinstance(item, dict):
            continue

        item_type = item.get("type", "text")
        if item_type == "text":
            metrics["text_blocks"] += 1
            text = item.get("text", "") or ""
            if not str(text).strip():
                metrics["empty_text_blocks"] += 1
            else:
                text_parts.append(str(text))
                metrics["text_chars"] += len(str(text))
        elif item_type in {"image", "chart"}:
            metrics["image_blocks"] += 1
            if item_type == "image":
                img_path = item.get("img_path")
                if img_path:
                    try:
                        import os

                        if os.path.exists(img_path) and _is_valid_image_file(img_path):
                            metrics["image_files_exist"] += 1
                        else:
                            metrics["image_files_missing"] += 1
                    except Exception:
                        metrics["image_files_missing"] += 1
                else:
                    metrics["image_files_missing"] += 1
        elif item_type == "table":
            metrics["table_blocks"] += 1
            table_body = item.get("table_body")
            tb = _parse_table_body(table_body)
            metrics["table_rows"] += tb["rows"]
            metrics["table_cells"] += tb["cells"]
        elif item_type == "equation":
            metrics["equation_blocks"] += 1

    joined_text = "\n\n".join(text_parts)
    metrics["text_tokens"] = _count_tokens(joined_text)
    metrics["text_md5"] = hashlib.md5(joined_text.encode("utf-8", errors="ignore")).hexdigest()
    metrics["noise_ratio"] = _estimate_noise_ratio(content_list, metrics["text_chars"])

    pages = metrics["source_pages"]
    metrics["tokens_per_page"] = _safe_rate(metrics["text_tokens"], pages)
    metrics["text_blocks_per_page"] = _safe_rate(metrics["text_blocks"], pages)
    metrics["tables_per_100_pages"] = _safe_rate(metrics["table_blocks"], pages, scale=100.0)
    metrics["figures_per_100_pages"] = _safe_rate(metrics["image_blocks"], pages, scale=100.0)
    metrics["equations_per_100_pages"] = _safe_rate(metrics["equation_blocks"], pages, scale=100.0)
    metrics["modality_coverage_profile"] = _format_modality_profile(metrics)

    return metrics


def summarize_extract_metrics(rows: Iterable[Dict[str, Any]]) -> Dict[str, Any]:
    rows = list(rows)
    total_files = len(rows)
    success_rows = [row for row in rows if row.get("status") == "Success"]
    success_files = len(success_rows)

    def _median(key: str) -> float:
        values = [float(row.get(key, 0.0) or 0.0) for row in success_rows]
        return float(median(values)) if values else 0.0

    def _mean(key: str) -> float:
        values = [float(row.get(key, 0.0) or 0.0) for row in success_rows]
        return float(sum(values) / len(values)) if values else 0.0

    summary = {
        "files_total": total_files,
        "files_succeeded": success_files,
        "parse_success_rate": (success_files / total_files) if total_files else 0.0,
        "median_seconds_per_page": _median("seconds_per_page"),
        "median_noise_ratio": _median("noise_ratio"),
        "median_tokens_per_page": _median("tokens_per_page"),
        "mean_text_blocks_per_page": _mean("text_blocks_per_page"),
        "mean_tables_per_100_pages": _mean("tables_per_100_pages"),
        "mean_figures_per_100_pages": _mean("figures_per_100_pages"),
        "mean_equations_per_100_pages": _mean("equations_per_100_pages"),
    }
    summary["modality_coverage_profile"] = (
        f"text={summary['mean_text_blocks_per_page']:.1f}/page | "
        f"table={summary['mean_tables_per_100_pages']:.1f}/100p | "
        f"figure={summary['mean_figures_per_100_pages']:.1f}/100p | "
        f"equation={summary['mean_equations_per_100_pages']:.1f}/100p"
    )
    return summary
